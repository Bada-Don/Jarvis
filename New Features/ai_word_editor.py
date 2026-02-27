import os
import difflib
import subprocess
import platform
import docx
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize OpenAI client with API key from .env
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- 1. Structured Output Schema ---
class EditCommand(BaseModel):
    search_text: str
    replace_text: str

class FileEdits(BaseModel):
    edits: list[EditCommand]

# --- 2. Word Doc Helper Functions ---
def extract_text_from_docx(doc: docx.Document) -> str:
    """Extracts plain text from a Word document (paragraphs and tables)."""
    text_blocks = []
    
    # Extract standard paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_blocks.append(p.text)
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_blocks.append(p.text)
                        
    return "\n".join(text_blocks)

def apply_docx_edits(doc: docx.Document, edits: list[EditCommand]) -> bool:
    """
    Applies search/replace while aggressively preserving formatting (Font Size, Bold, Color, etc.)
    """
    changes_made = False
    
    # Gather all paragraphs
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
                
    for edit in edits:
        found = False
        for p in all_paragraphs:
            if edit.search_text in p.text:
                found = True
                changes_made = True
                
                # STRATEGY A: Try to replace text perfectly inside a single Run
                replaced_in_run = False
                for run in p.runs:
                    if edit.search_text in run.text:
                        run.text = run.text.replace(edit.search_text, edit.replace_text)
                        replaced_in_run = True
                
                # STRATEGY B: Text spans multiple runs. We must rebuild the paragraph,
                # but we will steal the formatting from the original text first.
                if not replaced_in_run:
                    # 1. Find the first run that actually has text to use as our formatting template
                    ref_run = p.runs[0] if p.runs else None
                    for run in p.runs:
                        if run.text.strip():
                            ref_run = run
                            break
                    
                    # 2. Extract formatting attributes safely
                    f_name = ref_run.font.name if ref_run else None
                    f_size = ref_run.font.size if ref_run else None
                    f_bold = ref_run.font.bold if ref_run else None
                    f_italic = ref_run.font.italic if ref_run else None
                    f_underline = ref_run.font.underline if ref_run else None
                    f_color = None
                    if ref_run and ref_run.font.color and ref_run.font.color.type == 1: # 1 is RGB
                        f_color = ref_run.font.color.rgb

                    # 3. Perform the text replacement on the plain text
                    new_text = p.text.replace(edit.search_text, edit.replace_text)
                    
                    # 4. Clear the paragraph (deletes all old runs)
                    p.clear()
                    
                    # 5. Rebuild the paragraph with the saved formatting
                    new_run = p.add_run(new_text)
                    if f_name: new_run.font.name = f_name
                    if f_size: new_run.font.size = f_size
                    if f_bold is not None: new_run.font.bold = f_bold
                    if f_italic is not None: new_run.font.italic = f_italic
                    if f_underline is not None: new_run.font.underline = f_underline
                    if f_color: new_run.font.color.rgb = f_color

        if not found:
            print(f"\n[WARNING] Could not find exact text to replace:\n'''{edit.search_text}'''\n")
            
    return changes_made

# --- 3. Core AI Function ---
def get_ai_edits(file_content: str, prompt: str) -> FileEdits:
    system_prompt = """
    You are an expert file editor. Your task is to modify the provided text content based on the user's prompt.
    You must output a list of search and replace operations.
    CRITICAL RULES:
    1. 'search_text' must be an EXACT match of the text currently in the file.
    2. 'replace_text' is the new text that will replace 'search_text'.
    3. TO PRESERVE FORMATTING: Keep 'search_text' as SHORT and specific as possible. Instead of replacing a whole paragraph, just replace the exact words/names/values that need changing.
    4. Do NOT make 'search_text' span across multiple paragraphs.
    """

    completion = client.beta.chat.completions.parse(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Document Text:\n```\n{file_content}\n```\n\nUser Request: {prompt}"}
        ],
        response_format=FileEdits,
    )

    return completion.choices[0].message.parsed

# --- 4. Utilities ---
def show_diff(original_text: str, new_text: str, filepath: str):
    """Generates a unified diff using the extracted plain text."""
    diff = difflib.unified_diff(
        original_text.splitlines(keepends=True),
        new_text.splitlines(keepends=True),
        fromfile=f"a/{os.path.basename(filepath)} (Plain Text)",
        tofile=f"b/{os.path.basename(filepath)} (Plain Text)",
        n=3
    )
    diff_text = "".join(diff)
    if diff_text:
        print("\n--- PROPOSED CHANGES (TEXT DIFF) ---")
        print(diff_text)
        print("------------------------------------\n")

def open_file_in_editor(filepath: str):
    """Opens the file using the default OS viewer (e.g., MS Word)."""
    if platform.system() == 'Darwin':
        subprocess.call(('open', filepath))
    elif platform.system() == 'Windows':
        os.startfile(filepath)
    else:
        subprocess.call(('xdg-open', filepath))

# --- 5. Main Workflow ---
def main():
    print("🤖 AI Word Document Editor Agent")
    
    # Step 1: Get File Path
    filepath = input("1. Enter the path to the .docx file: ").strip()
    
    if not os.path.isfile(filepath) or not filepath.endswith('.docx'):
        print(f"Error: Valid .docx file not found at '{filepath}'.")
        return

    # Load Document and Extract Text
    doc = docx.Document(filepath)
    original_text = extract_text_from_docx(doc)

    # Step 2: Get User Prompt
    prompt = input('2. What would you like to do? (e.g., "Replace placeholders with Harshit Singla"): \n> ').strip()

    print("\nThinking and generating edits...")
    
    # Step 3: AI Edits
    try:
        parsed_edits = get_ai_edits(original_text, prompt)
    except Exception as e:
        print(f"Failed to communicate with AI: {e}")
        return

    # Apply edits to the document object in memory
    changes_made = apply_docx_edits(doc, parsed_edits.edits)
    
    if not changes_made:
        print("No valid changes were applied. Exiting.")
        return

    # Extract the new text to create a diff
    new_text = extract_text_from_docx(doc)
    show_diff(original_text, new_text, filepath)

    # Step 4: Confirm and Apply Changes
    confirm = input("4. Apply these changes and save? (y/n): ").strip().lower()
    if confirm == 'y':
        # Save the modified document object back to the file
        doc.save(filepath)
        print(f"Success! Changes saved to {filepath}.")
        
        # Step 5: Open the file
        print("5. Opening MS Word...")
        open_file_in_editor(filepath)
    else:
        print("Changes discarded.")

if __name__ == "__main__":
    main()