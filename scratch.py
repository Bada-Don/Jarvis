import sys
from pathlib import Path
backend_path = Path(__file__).resolve().parent / "backend"
sys.path.insert(0, str(backend_path))

from ai_editor_engine import AIEditorEngine, WordEditCommand
import docx

doc = docx.Document()
engine = AIEditorEngine(api_key="dummy")
edits = [WordEditCommand(search_text="", replace_text="AIM: To write a program\n\nCode:\nimport math")]

changed = engine.apply_word_edits(doc, edits)
print(f"Changes made: {changed}")
doc.save('test_engine.docx')

doc2 = docx.Document('test_engine.docx')
print("Read back:")
for p in doc2.paragraphs:
    print(repr(p.text))
