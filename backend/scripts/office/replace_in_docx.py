"""
scripts/office/replace_in_docx.py
Find-and-replace text in a .docx file, preserving all formatting.

Usage:
    python scripts/office/replace_in_docx.py input.docx output.docx "old text" "new text"
    python scripts/office/replace_in_docx.py input.docx output.docx "old" "new" --case-sensitive false

Notes:
  - Replaces in all paragraph runs AND table cell paragraphs, headers, footers
  - Handles fragmented runs: collapses adjacent runs with identical rPr before matching
  - Preserves bold, italic, font, color, and all other run-level formatting
  - Reports total replacement count on completion
"""

import argparse
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ── Run merging ──────────────────────────────────────────────────────────────

def _runs_have_same_format(r1, r2) -> bool:
    rpr1 = r1._r.find(qn("w:rPr"))
    rpr2 = r2._r.find(qn("w:rPr"))
    if rpr1 is None and rpr2 is None:
        return True
    if rpr1 is None or rpr2 is None:
        return False
    return etree.tostring(rpr1) == etree.tostring(rpr2)


def _merge_paragraph_runs(para) -> None:
    """Merge adjacent runs with identical rPr in-place."""
    i = 0
    while True:
        runs = para.runs
        if i >= len(runs) - 1:
            break
        r1, r2 = runs[i], runs[i + 1]
        if _runs_have_same_format(r1, r2):
            r1.text = (r1.text or "") + (r2.text or "")
            r2._r.getparent().remove(r2._r)
            # Don't advance i — re-check at same position
        else:
            i += 1


# ── Replace in paragraph ─────────────────────────────────────────────────────

def _replace_in_paragraph(para, old: str, new: str, case_sensitive: bool) -> int:
    """Merge runs, then replace. Returns replacement count."""
    _merge_paragraph_runs(para)
    count = 0
    for run in para.runs:
        text = run.text or ""
        if case_sensitive:
            if old in text:
                count += text.count(old)
                run.text = text.replace(old, new)
        else:
            pattern = re.compile(re.escape(old), re.IGNORECASE)
            matches = pattern.findall(text)
            if matches:
                count += len(matches)
                run.text = pattern.sub(new, text)
    return count


# ── Public API ───────────────────────────────────────────────────────────────

def replace_in_docx(
    src: str,
    dst: str,
    old: str,
    new: str,
    case_sensitive: bool = True,
) -> int:
    """
    Open `src`, replace all occurrences of `old` with `new`, save to `dst`.
    Covers: body paragraphs, table cells, headers, footers.
    Returns total replacement count.
    """
    doc = Document(src)
    total = 0

    # Body paragraphs
    for para in doc.paragraphs:
        total += _replace_in_paragraph(para, old, new, case_sensitive)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += _replace_in_paragraph(para, old, new, case_sensitive)

    # Headers and footers across all sections
    for section in doc.sections:
        for hf in [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ]:
            if hf is not None:
                for para in hf.paragraphs:
                    total += _replace_in_paragraph(para, old, new, case_sensitive)

    doc.save(dst)
    return total


# ── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Find-and-replace text in a .docx file"
    )
    parser.add_argument("src", help="Input .docx file")
    parser.add_argument("dst", help="Output .docx file")
    parser.add_argument("old", help="Text to find")
    parser.add_argument("new", help="Replacement text")
    parser.add_argument(
        "--case-sensitive",
        default="true",
        choices=["true", "false"],
        help="Case-sensitive match (default: true)",
    )
    args = parser.parse_args()

    n = replace_in_docx(
        args.src, args.dst, args.old, args.new,
        case_sensitive=(args.case_sensitive == "true"),
    )

    if n == 0:
        print(f"⚠️  No occurrences of '{args.old}' found in '{args.src}'")
        sys.exit(1)
    else:
        print(f"✅ Replaced {n} occurrence(s): '{args.old}' → '{args.new}'")
        print(f"   Saved: '{args.dst}'")


if __name__ == "__main__":
    main()
