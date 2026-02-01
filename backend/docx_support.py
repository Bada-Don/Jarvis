"""
Optional DOCX Support for Jarvis File Operations

Install with: pip install python-docx

This module extends file_operations.py to support reading Word documents.
"""

from typing import Tuple, Optional
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")


def read_docx(path: str) -> Tuple[bool, str, Optional[str]]:
    """
    Read text content from a .docx Word document.
    
    Args:
        path: Absolute or relative file path to .docx file
    
    Returns:
        Tuple[bool, str, Optional[str]]: (success, message, content)
        If success is False, content will be None
    
    Example:
        >>> success, msg, content = read_docx("~/Desktop/AI Lab/Practical 1.docx")
        >>> if success:
        ...     print(content)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx library not installed. Install with: pip install python-docx", None
    
    try:
        file_path = Path(path).expanduser()
        
        if not file_path.exists():
            return False, f"File not found: {path}", None
        
        if not file_path.is_file():
            return False, f"Path is not a file: {path}", None
        
        if not str(file_path).lower().endswith('.docx'):
            return False, f"File is not a .docx document: {path}", None
        
        # Read the document
        doc = Document(file_path)
        
        # Extract all text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content = '\n'.join(paragraphs)
        
        # Also extract text from tables if present
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        if tables_text:
            content += '\n\n' + '\n'.join(tables_text)
        
        return True, f"Document read successfully: {file_path}", content
        
    except PermissionError:
        return False, f"Permission denied: {path}", None
    except Exception as e:
        return False, f"Error reading document: {e}", None


def read_file_smart(path: str, encoding: str = 'utf-8') -> Tuple[bool, str, Optional[str]]:
    """
    Smart file reader that automatically detects file type.
    Supports both text files and .docx documents.
    
    Args:
        path: Absolute or relative file path
        encoding: File encoding for text files (default: utf-8)
    
    Returns:
        Tuple[bool, str, Optional[str]]: (success, message, content)
    """
    file_path = Path(path).expanduser()
    
    # Check if it's a .docx file
    if str(file_path).lower().endswith('.docx'):
        return read_docx(path)
    
    # Otherwise, use standard text file reading
    try:
        if not file_path.exists():
            return False, f"File not found: {path}", None
        
        if not file_path.is_file():
            return False, f"Path is not a file: {path}", None
        
        content = file_path.read_text(encoding=encoding)
        return True, f"File read successfully: {file_path}", content
        
    except PermissionError:
        return False, f"Permission denied: {path}", None
    except UnicodeDecodeError:
        return False, f"Encoding error (try different encoding): {path}", None
    except Exception as e:
        return False, f"Error reading file: {e}", None


# Example usage
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_support.py <path_to_docx_file>")
        print("\nExample:")
        print('  python docx_support.py "C:\\Users\\user\\Desktop\\AI Lab\\Practical 1.docx"')
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    print(f"Reading: {file_path}")
    print("=" * 60)
    
    success, message, content = read_docx(file_path)
    
    print(f"Status: {message}")
    
    if success and content:
        print("\nContent:")
        print("-" * 60)
        print(content)
        print("-" * 60)
        print(f"\nTotal characters: {len(content)}")
        print(f"Total lines: {len(content.splitlines())}")
    else:
        print(f"\nError: {message}")
