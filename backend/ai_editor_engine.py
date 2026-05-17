import os
import difflib
import openpyxl
from openpyxl.utils import get_column_letter
import docx
from pydantic import BaseModel
from typing import Literal, List, Optional, Union
from google import genai
from google.genai import types
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# --- 1. Schemas ---

class TextEditCommand(BaseModel):
    search_text: str
    replace_text: str

class TextFileEdits(BaseModel):
    edits: List[TextEditCommand]

class ExcelCommand(BaseModel):
    action: Literal["edit_cell", "insert_row", "delete_row"]
    sheet_name: str
    target: str  # For edits: "A7". For inserts/deletes: "7"
    value: str   # The new value (leave empty for inserts/deletes)

class ExcelFileEdits(BaseModel):
    commands: List[ExcelCommand]

class WordEditCommand(BaseModel):
    search_text: str
    replace_text: str

class WordFileEdits(BaseModel):
    edits: List[WordEditCommand]

# Type alias for optional type hints
AIEditorEngineType = Optional['AIEditorEngine']

# --- 2. AI Editor Engine Class ---

class AIEditorEngine:
    def __init__(self, api_key: Optional[str] = None):
        api_key = api_key or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment or provided.")
        self.client = genai.Client(api_key=api_key)
        self.model = "gemini-2.5-flash"
        self.last_usage: Optional[dict] = None

    def _update_usage(self, response):
        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            self.last_usage = {
                'prompt_tokens': response.usage_metadata.prompt_token_count,
                'candidates_tokens': response.usage_metadata.candidates_token_count,
                'total_tokens': response.usage_metadata.total_token_count,
            }
            if hasattr(response.usage_metadata, 'thoughts_token_count'):
                self.last_usage['thoughts_tokens'] = response.usage_metadata.thoughts_token_count

    # --- Text Editing ---
    def get_text_edits(self, content: str, prompt: str) -> TextFileEdits:
        system_prompt = """
        You are an expert coding assistant and file editor. 
        Your task is to modify the provided file content based on the user's prompt.
        You must output a list of search and replace operations.
        CRITICAL RULES:
        1. 'search_text' must be an EXACT match of the text currently in the file. Include exact whitespace and indentation.
        2. 'replace_text' is the new text that will replace 'search_text'.
        3. Make the 'search_text' uniquely identifiable (include a few lines above/below if necessary).
        """
        response = self.client.models.generate_content(
            model=self.model,
            contents=f"File Content:\n```\n{content}\n```\n\nUser Request: {prompt}",
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type='application/json',
                response_schema=TextFileEdits,
            ),
        )
        self._update_usage(response)
        return response.parsed

    def apply_text_edits(self, content: str, edits: List[TextEditCommand]) -> str:
        new_content = content
        for edit in edits:
            if edit.search_text in new_content:
                new_content = new_content.replace(edit.search_text, edit.replace_text)
        return new_content

    # --- Excel Editing ---
    def extract_excel_context(self, wb: openpyxl.Workbook, max_rows: int = 100) -> str:
        context = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            context.append(f"--- SHEET: {sheet_name} ---")
            max_col = ws.max_column
            col_headers = ["ROW/COL"] + [get_column_letter(i) for i in range(1, max_col + 1)]
            context.append(" | ".join(col_headers))
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if row_idx > max_rows:
                    context.append(f"... (Truncated at {max_rows} rows) ...")
                    break
                row_data = [str(row_idx)] + [str(cell_val).replace("\n", " ") if cell_val is not None else "" for cell_val in row]
                context.append(" | ".join(row_data))
            context.append("\n")
        return "\n".join(context)

    def get_excel_edits(self, context: str, prompt: str) -> ExcelFileEdits:
        system_prompt = """
        You are an expert Data Analyst and Excel editor. 
        You have been provided with a text representation of an Excel workbook.
        
        CRITICAL RULES:
        1. ACTIONS: You can 'insert_row', 'delete_row', or 'edit_cell'. 
        2. SEQUENTIAL LOGIC: Commands are executed in the exact order you provide.
        3. TARGETING: For row operations, the 'target' is the row number (e.g., '7'). For cell edits, the target is the coordinate (e.g., 'A7').
        4. ADDING DATA: If asked to add a new row between existing data, first use 'insert_row', then follow up with multiple 'edit_cell' commands.
        5. DATA TYPES: Output just the numbers if it's a number. Output exact formulas starting with '=' if requested.
        """
        response = self.client.models.generate_content(
            model=self.model,
            contents=f"Excel Data:\n```text\n{context}\n```\n\nUser Request: {prompt}",
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type='application/json',
                response_schema=ExcelFileEdits,
            ),
        )
        self._update_usage(response)
        return response.parsed

    def apply_excel_edits(self, wb: openpyxl.Workbook, commands: List[ExcelCommand]) -> List[dict]:
        diff_records = []
        for cmd in commands:
            if cmd.sheet_name not in wb.sheetnames: continue
            ws = wb[cmd.sheet_name]
            try:
                if cmd.action == "insert_row":
                    ws.insert_rows(int(cmd.target))
                    diff_records.append({"action": "insert", "sheet": cmd.sheet_name, "target": cmd.target})
                elif cmd.action == "delete_row":
                    ws.delete_rows(int(cmd.target))
                    diff_records.append({"action": "delete", "sheet": cmd.sheet_name, "target": cmd.target})
                elif cmd.action == "edit_cell":
                    target_cell = ws[cmd.target]
                    old_val = str(target_cell.value) if target_cell.value is not None else "(empty)"
                    parsed_val = self._parse_excel_value(cmd.value)
                    target_cell.value = parsed_val
                    diff_records.append({
                        "action": "edit", "sheet": cmd.sheet_name, "target": cmd.target,
                        "old": old_val, "new": cmd.value, "type": type(parsed_val).__name__
                    })
            except Exception: continue
        return diff_records

    def _parse_excel_value(self, val_str: str):
        val_str = val_str.strip()
        if not val_str: return ""
        if val_str.startswith("="): return val_str
        try: return int(val_str)
        except ValueError: pass
        try: return float(val_str)
        except ValueError: pass
        if val_str.lower() == "true": return True
        if val_str.lower() == "false": return False
        return val_str

    # --- Word Editing ---
    def extract_word_context(self, doc: docx.Document) -> str:
        text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_blocks.extend([p.text for p in cell.paragraphs if p.text.strip()])
        return "\n".join(text_blocks)

    def get_word_edits(self, context: str, prompt: str) -> WordFileEdits:
        system_prompt = """
        You are an expert Word document editor. Your task is to modify the provided text content based on the user's prompt.
        You must output a list of search and replace operations.
        CRITICAL RULES:
        1. 'search_text' must be an EXACT match of the text currently in the file.
        2. 'replace_text' is the new text that will replace 'search_text'.
        3. TO PRESERVE FORMATTING: Keep 'search_text' as SHORT and specific as possible.
        4. Do NOT make 'search_text' span across multiple paragraphs.
        """
        response = self.client.models.generate_content(
            model=self.model,
            contents=f"Document Text:\n```\n{context}\n```\n\nUser Request: {prompt}",
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type='application/json',
                response_schema=WordFileEdits,
            ),
        )
        self._update_usage(response)
        return response.parsed

    def apply_word_edits(self, doc: docx.Document, edits: List[WordEditCommand]) -> bool:
        changes_made = False
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for edit in edits:
            if not edit.search_text:
                if edit.replace_text:
                    # If document is virtually empty, use the first paragraph
                    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
                        doc.paragraphs[0].text = edit.replace_text
                    else:
                        doc.add_paragraph(edit.replace_text)
                    changes_made = True
                continue

            # Search and replace
            found = False
            for p in all_paragraphs:
                if edit.search_text in p.text:
                    # To preserve formatting, we try to replace within runs if possible
                    replaced_in_run = False
                    for run in p.runs:
                        if edit.search_text in run.text:
                            run.text = run.text.replace(edit.search_text, edit.replace_text)
                            replaced_in_run = True
                            changes_made = True
                            found = True
                    
                    if not replaced_in_run:
                        p.text = p.text.replace(edit.search_text, edit.replace_text)
                        changes_made = True
                        found = True

            # If not found and it's a new/empty document, just add the text
            if not found and edit.replace_text:
                if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
                    doc.paragraphs[0].text = edit.replace_text
                    changes_made = True
                else:
                    # Robustness: if document was supposed to be new but has content, append
                    pass

        return changes_made

    # --- Diff Generator ---
    def generate_text_diff(self, original: str, new: str, filename: str) -> str:
        diff = difflib.unified_diff(
            original.splitlines(keepends=True),
            new.splitlines(keepends=True),
            fromfile=f"a/{filename}", tofile=f"b/{filename}", n=3
        )
        return "".join(diff)

    def generate_excel_diff_summary(self, diff_records: List[dict]) -> str:
        lines = []
        for rec in diff_records:
            if rec['action'] == 'insert': lines.append(f"[{rec['sheet']}] ➕ Inserted row at {rec['target']}")
            elif rec['action'] == 'delete': lines.append(f"[{rec['sheet']}] ❌ Deleted row {rec['target']}")
            elif rec['action'] == 'edit': lines.append(f"[{rec['sheet']}] ✏️ {rec['target']}: '{rec['old']}' -> '{rec['new']}'")
        return "\n".join(lines)
